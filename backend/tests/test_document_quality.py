from pathlib import Path

import fitz
from docx import Document

from app.core.config import settings
from app.services.documents import (
    clean_page_text,
    extract_blocks,
    extract_pages,
    is_short_noise,
    markdown_table,
    repeated_margin_lines,
    response_text,
    structured_split,
)


def test_repeated_headers_footers_and_page_numbers_are_removed():
    pages = [
        (1, "企业资料汇编 | 内部测试\n1. 公司信息\n正文甲\n第 1 页"),
        (2, "企业资料汇编 | 内部测试\n2. 服务政策\n正文乙\n第 2 页"),
        (3, "企业资料汇编 | 内部测试\n3. 联系方式\n正文丙\n第 3 页"),
    ]
    repeated = repeated_margin_lines(pages)

    cleaned = clean_page_text(pages[0][1], repeated)

    assert "企业资料汇编" not in cleaned
    assert "第 1 页" not in cleaned
    assert "正文甲" in cleaned


def test_dynamic_date_page_header_is_removed():
    cleaned = clean_page_text(
        "资料采集日期：2026年6月13日 | 第 3 页\n3. 商城政策\n正文内容",
        set(),
    )

    assert "资料采集日期" not in cleaned
    assert "正文内容" in cleaned


def test_short_business_document_is_not_noise():
    assert not is_short_noise("核心卖点是部署简单，回答带有引用。")
    assert is_short_noise("来源：https://example.com")


def test_structured_split_preserves_section_metadata_and_overlap():
    text = "1. 公司信息\n" + "华为提供企业服务。" * 80 + "\n2. 服务政策\n支持预约维修和寄修。"

    chunks = structured_split(text, size=120, overlap=20)

    assert len(chunks) > 2
    assert chunks[0].section_title == "1. 公司信息"
    assert chunks[-1].section_title == "2. 服务政策"
    assert chunks[0].char_end > chunks[0].char_start


def test_extract_docx_preserves_heading_and_body(tmp_path: Path):
    path = tmp_path / "business.docx"
    document = Document()
    document.add_heading("订单管理", level=1)
    document.add_paragraph("订单提交后不能修改地址。")
    document.save(path)

    pages = extract_pages(path)

    assert pages == [(None, "# 订单管理\n订单提交后不能修改地址。")]


def test_extract_csv_and_html(tmp_path: Path):
    csv_path = tmp_path / "products.csv"
    html_path = tmp_path / "faq.html"
    csv_path.write_text("系列,定位\nMate,商务", encoding="utf-8")
    html_path.write_text("<h1>支付方式</h1><p>支持银行卡。</p>", encoding="utf-8")

    assert "Mate | 商务" in extract_pages(csv_path)[0][1]
    assert "支付方式" in extract_pages(html_path)[0][1]
    assert "支持银行卡" in extract_pages(html_path)[0][1]


def test_markdown_table_preserves_complex_cells():
    content = markdown_table([["产品", "规格"], ["Mate", "内存 12 GB\n存储 512 GB"]])

    assert "| 产品 | 规格 |" in content
    assert "| Mate | 内存 12 GB 存储 512 GB |" in content
    assert "| --- | --- |" in content


def test_extract_docx_tables_as_separate_blocks(tmp_path: Path):
    path = tmp_path / "products.docx"
    document = Document()
    document.add_paragraph("产品参数")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "型号"
    table.cell(0, 1).text = "容量"
    table.cell(1, 0).text = "Mate"
    table.cell(1, 1).text = "512 GB"
    document.save(path)

    blocks = extract_blocks(path)

    assert any(block.content_type == "body" and "产品参数" in block.content for block in blocks)
    assert any(block.content_type == "table" and "512 GB" in block.content for block in blocks)


def test_scanned_pdf_uses_vision_ocr(tmp_path: Path, monkeypatch):
    path = tmp_path / "scan.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(path)
    monkeypatch.setattr(settings, "ocr_min_page_characters", 80)
    monkeypatch.setattr(
        "app.services.documents.analyze_image",
        lambda *_args: "扫描识别结果\n| 项目 | 数值 |\n| --- | --- |\n| 收入 | 100 万 |",
    )

    blocks = extract_blocks(path)

    assert any(block.content_type == "ocr" and "扫描识别结果" in block.content for block in blocks)


def test_standalone_image_uses_vision_analysis(tmp_path: Path, monkeypatch):
    path = tmp_path / "chart.png"
    path.write_bytes(b"\x89PNG\r\n\x1a\nfake")
    monkeypatch.setattr(
        "app.services.documents.analyze_image",
        lambda *_args: "图表显示企业业务收入增长 20%。",
    )

    blocks = extract_blocks(path)

    assert blocks[0].content_type == "image"
    assert "增长 20%" in blocks[0].content


def test_vision_response_text_supports_multimodal_content_blocks():
    content = [{"type": "text", "text": "第一段"}, {"type": "text", "text": "第二段"}]

    assert response_text(content) == "第一段\n第二段"
