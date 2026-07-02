import base64
import json
import mimetypes
import re
from collections import Counter
from csv import reader
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

import fitz
import httpx
from docx import Document as WordDocument
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import ChunkEmbedding, Document, DocumentChunk
from app.services.rag import embed_text

HEADING_PATTERN = re.compile(
    r"^(?:#{1,6}\s+.+|\d+(?:\.\d+){0,3}\.?\s+.+|[一二三四五六七八九十]+[、.]\s*.+)$"
)
PAGE_NUMBER_PATTERN = re.compile(r"^(?:第\s*\d+\s*页|\d+\s*/\s*\d+|\d+)$")
TOC_PATTERN = re.compile(r".+\.{3,}\s*\d+$")


class TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "br", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", "".join(self.parts)).strip()


@dataclass
class StructuredChunk:
    content: str
    section_title: str
    char_start: int
    char_end: int
    content_type: str = "body"


@dataclass
class ExtractedBlock:
    page_number: int | None
    content: str
    content_type: str = "body"
    section_title: str = ""


def markdown_table(rows: list[list[str | None]]) -> str:
    normalized = [[normalized_line(cell or "") for cell in row] for row in rows]
    normalized = [row for row in normalized if any(row)]
    if not normalized:
        return ""
    width = max(len(row) for row in normalized)
    padded = [row + [""] * (width - len(row)) for row in normalized]
    header = padded[0]
    separator = ["---"] * width
    return "\n".join(
        "| " + " | ".join(row) + " |" for row in [header, separator, *padded[1:]]
    )


def vision_endpoint(base_url: str) -> str:
    return f"{base_url.rstrip('/')}/chat/completions"


def response_text(content: str | list[dict]) -> str:
    if isinstance(content, str):
        return content.strip()
    return "\n".join(
        item.get("text", "").strip()
        for item in content
        if isinstance(item, dict) and item.get("type") == "text"
    ).strip()


def analyze_image(image_bytes: bytes, mime_type: str, prompt: str) -> str:
    if not settings.vision_enabled:
        return ""
    api_key = settings.vision_api_key or settings.llm_api_key
    base_url = settings.vision_base_url or settings.llm_base_url
    model = settings.vision_model or settings.llm_model
    if not api_key or not base_url or not model:
        return ""
    encoded = base64.b64encode(image_bytes).decode("ascii")
    try:
        response = httpx.post(
            vision_endpoint(base_url),
            headers={"Authorization": f"Bearer {api_key}"},
            json={
                "model": model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:{mime_type};base64,{encoded}"},
                            },
                        ],
                    }
                ],
                "temperature": 0,
            },
            timeout=settings.vision_timeout_seconds,
        )
        response.raise_for_status()
        return response_text(response.json()["choices"][0]["message"]["content"])
    except (httpx.HTTPError, KeyError, TypeError, ValueError):
        return ""


def extract_pdf_blocks(path: Path) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    analyzed_images = 0
    with fitz.open(path) as pdf:
        for index, page in enumerate(pdf):
            page_number = index + 1
            text = page.get_text().strip()
            if text:
                blocks.append(ExtractedBlock(page_number, text))

            try:
                tables = page.find_tables()
                for table_index, table in enumerate(tables.tables, start=1):
                    content = markdown_table(table.extract())
                    if content:
                        blocks.append(
                            ExtractedBlock(
                                page_number,
                                content,
                                "table",
                                f"第 {page_number} 页表格 {table_index}",
                            )
                        )
            except (AttributeError, RuntimeError, ValueError):
                pass

            if len(normalized_line(text)) < settings.ocr_min_page_characters:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                ocr_text = analyze_image(
                    pixmap.tobytes("png"),
                    "image/png",
                    "请完整识别此扫描文档页。按原阅读顺序输出全部文字；表格请使用 Markdown "
                    "表格；不要解释，不要省略数字。",
                )
                if ocr_text:
                    blocks.append(ExtractedBlock(page_number, ocr_text, "ocr", "扫描页 OCR"))
                continue

            for image_index, image in enumerate(page.get_images(full=True), start=1):
                if analyzed_images >= settings.vision_max_images_per_document:
                    break
                if image[2] < 128 or image[3] < 128:
                    continue
                extracted = pdf.extract_image(image[0])
                extension = extracted.get("ext", "png")
                mime_type = "image/jpeg" if extension in {"jpg", "jpeg"} else f"image/{extension}"
                description = analyze_image(
                    extracted["image"],
                    mime_type,
                    "请描述这张文档图片表达的信息，并提取图片中的文字、数字、图例和关键结论。"
                    "只输出可用于知识库检索的事实，不要猜测。",
                )
                if description:
                    blocks.append(
                        ExtractedBlock(
                            page_number,
                            description,
                            "image",
                            f"第 {page_number} 页图片 {image_index}",
                        )
                    )
                analyzed_images += 1
    return blocks


def extract_docx_blocks(path: Path) -> list[ExtractedBlock]:
    document = WordDocument(path)
    blocks: list[ExtractedBlock] = []
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = normalized_line(paragraph.text)
        if not text:
            continue
        lines.append(f"# {text}" if paragraph.style.name.startswith("Heading") else text)
    if lines:
        blocks.append(ExtractedBlock(None, "\n".join(lines)))
    for table_index, table in enumerate(document.tables, start=1):
        content = markdown_table([[cell.text for cell in row.cells] for row in table.rows])
        if content:
            blocks.append(ExtractedBlock(None, content, "table", f"表格 {table_index}"))
    analyzed_images = 0
    for part in document.part.related_parts.values():
        if not getattr(part, "content_type", "").startswith("image/"):
            continue
        if analyzed_images >= settings.vision_max_images_per_document:
            break
        description = analyze_image(
            part.blob,
            part.content_type,
            "请描述这张文档图片表达的信息，并提取图片中的文字、数字、图例和关键结论。"
            "只输出可用于知识库检索的事实，不要猜测。",
        )
        if description:
            blocks.append(
                ExtractedBlock(None, description, "image", f"文档图片 {analyzed_images + 1}")
            )
        analyzed_images += 1
    return blocks


def extract_blocks(path: Path) -> list[ExtractedBlock]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_blocks(path)
    if suffix == ".docx":
        return extract_docx_blocks(path)
    if suffix == ".csv":
        rows = reader(path.read_text(encoding="utf-8-sig", errors="ignore").splitlines())
        return [ExtractedBlock(None, markdown_table(list(rows)), "table", "CSV 表格")]
    if suffix in {".html", ".htm"}:
        parser = TextHTMLParser()
        parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
        return [ExtractedBlock(None, parser.text())]
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        mime_type = mimetypes.guess_type(path.name)[0] or "image/png"
        description = analyze_image(
            path.read_bytes(),
            mime_type,
            "请完整识别图片中的文字、表格、数字和视觉信息，并总结可验证的关键事实。"
            "表格请使用 Markdown 表格，不要猜测。",
        )
        return [ExtractedBlock(None, description, "image", "图片识别")] if description else []
    return [ExtractedBlock(None, path.read_text(encoding="utf-8", errors="ignore"))]


def extract_pages(path: Path) -> list[tuple[int | None, str]]:
    pages: dict[int | None, list[str]] = {}
    for block in extract_blocks(path):
        pages.setdefault(block.page_number, []).append(block.content)
    return [(page_number, "\n".join(parts)) for page_number, parts in pages.items()]


def normalized_line(line: str) -> str:
    return re.sub(r"\s+", " ", line).strip()


def repeated_margin_lines(pages: list[tuple[int | None, str]]) -> set[str]:
    if len(pages) < 2:
        return set()
    counts: Counter[str] = Counter()
    for _, text in pages:
        lines = [normalized_line(line) for line in text.splitlines() if normalized_line(line)]
        for line in set(lines[:3] + lines[-3:]):
            if 4 <= len(line) <= 160:
                counts[line] += 1
    threshold = max(2, int(len(pages) * 0.6))
    return {line for line, count in counts.items() if count >= threshold}


def clean_page_text(text: str, repeated_lines: set[str]) -> str:
    kept = []
    for raw_line in text.splitlines():
        line = normalized_line(raw_line)
        if (
            not line
            or line in repeated_lines
            or PAGE_NUMBER_PATTERN.fullmatch(line)
            or ("资料采集日期" in line and re.search(r"第\s*\d+\s*页", line))
        ):
            continue
        if TOC_PATTERN.fullmatch(line):
            continue
        kept.append(line)
    return "\n".join(kept).strip()


def is_heading(line: str) -> bool:
    return bool(HEADING_PATTERN.fullmatch(line)) and len(line) <= 100


def split_long_block(text: str, size: int, overlap: int) -> list[tuple[str, int, int]]:
    if len(text) <= size:
        return [(text, 0, len(text))]
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        if end < len(text):
            boundary = max(text.rfind("\n", start, end), text.rfind("。", start, end))
            if boundary > start + size // 2:
                end = boundary + 1
        chunks.append((text[start:end].strip(), start, end))
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


def structured_split(text: str, size: int = 700, overlap: int = 100) -> list[StructuredChunk]:
    cleaned = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not cleaned:
        return []
    sections: list[tuple[str, list[str]]] = []
    title = ""
    body: list[str] = []
    for line in cleaned.splitlines():
        if is_heading(line):
            if body:
                sections.append((title, body))
            title, body = line.lstrip("# ").strip(), []
        else:
            body.append(line)
    if body or title:
        sections.append((title, body))

    chunks: list[StructuredChunk] = []
    cursor = 0
    for section_title, lines in sections:
        section_text = "\n".join(lines).strip()
        if not section_text:
            continue
        prefix = f"{section_title}\n" if section_title else ""
        for content, local_start, local_end in split_long_block(section_text, size, overlap):
            full_content = f"{prefix}{content}".strip()
            chunks.append(
                StructuredChunk(
                    content=full_content,
                    section_title=section_title,
                    char_start=cursor + local_start,
                    char_end=cursor + local_end,
                )
            )
        cursor += len(section_text) + 1
    return chunks


def is_short_noise(content: str) -> bool:
    if len(content) >= 60:
        return False
    return (
        "资料采集日期" in content
        or content.startswith("来源：")
        or "官方来源目录" in content
        or bool(re.search(r"https?://", content))
    )


def split_text(text: str, size: int = 800, overlap: int = 120) -> list[str]:
    return [chunk.content for chunk in structured_split(text, size, overlap)]


def process_document(document_id: int, db: Session) -> None:
    document = db.get(Document, document_id)
    if not document or document.deleted_at is not None:
        raise ValueError("Document not found")
    document.status = "processing"
    document.error_message = ""
    document.chunks.clear()
    db.flush()
    blocks = extract_blocks(Path(document.file_path))
    body_pages: dict[int | None, list[str]] = {}
    for block in blocks:
        if block.content_type in {"body", "ocr"}:
            body_pages.setdefault(block.page_number, []).append(block.content)
    repeated_lines = repeated_margin_lines(
        [(page_number, "\n".join(parts)) for page_number, parts in body_pages.items()]
    )
    index = 0
    for block in blocks:
        cleaned = (
            clean_page_text(block.content, repeated_lines)
            if block.content_type in {"body", "ocr"}
            else block.content.strip()
        )
        for item in structured_split(cleaned, settings.chunk_size, settings.chunk_overlap):
            if is_short_noise(item.content):
                continue
            chunk = DocumentChunk(
                document_id=document.id,
                content=item.content,
                chunk_index=index,
                page_number=block.page_number,
                section_title=item.section_title or block.section_title,
                char_start=item.char_start,
                char_end=item.char_end,
                content_type=block.content_type,
            )
            db.add(chunk)
            db.flush()
            vector = embed_text(item.content)
            db.add(
                ChunkEmbedding(
                    chunk_id=chunk.id,
                    vector_json=json.dumps(vector),
                    vector=vector,
                    model=settings.embedding_model,
                )
            )
            index += 1
    if index == 0:
        raise ValueError("No readable text was extracted")
    document.status = "ready"
    db.commit()
